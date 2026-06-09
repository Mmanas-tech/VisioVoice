"""Document export service for DOCX and PDF formats."""

import io
import logging
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


def _format_timestamp_ms(ms: int) -> str:
    """Convert milliseconds to HH:MM:SS format."""
    hours = ms // 3600000
    minutes = (ms % 3600000) // 60000
    seconds = (ms % 60000) // 1000
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


class DocumentExporter:
    """Export transcriptions to DOCX and PDF formats."""

    def export_docx(
        self,
        segments: List[Dict[str, Any]],
        full_text: str,
        title: str = "Transcription",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> bytes:
        """Export transcription as DOCX document.

        Args:
            segments: List of segment dicts with start_ms, end_ms, text, confidence_score
            full_text: Full transcription text
            title: Document title
            metadata: Optional metadata dict (confidence, processing_time, etc.)

        Returns:
            DOCX file bytes
        """
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        doc.add_heading(title, level=0)

        if metadata:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            meta_run = meta_para.add_run("Metadata\n")
            meta_run.bold = True
            meta_run.font.size = Pt(10)
            meta_run.font.color.rgb = RGBColor(100, 100, 100)

            for key, value in metadata.items():
                p = doc.add_paragraph()
                r = p.add_run(f"{key}: ")
                r.bold = True
                r.font.size = Pt(10)
                r = p.add_run(str(value))
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(100, 100, 100)

            doc.add_paragraph()

        doc.add_heading("Full Transcription", level=1)
        doc.add_paragraph(full_text)

        if segments:
            doc.add_heading("Timestamped Segments", level=1)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Shading Accent 1"

            hdr = table.rows[0].cells
            hdr[0].text = "Time"
            hdr[1].text = "Segment"
            hdr[2].text = "Text"
            hdr[3].text = "Confidence"

            for cell in hdr:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for seg in segments:
                row = table.add_row().cells
                start = _format_timestamp_ms(int(seg.get("start_ms", 0)))
                end = _format_timestamp_ms(int(seg.get("end_ms", 0)))
                row[0].text = f"{start} → {end}"
                row[1].text = str(seg.get("segment_index", ""))
                row[2].text = seg.get("text", "")
                confidence = seg.get("confidence_score")
                row[3].text = f"{confidence:.1%}" if confidence is not None else "N/A"

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def export_pdf(
        self,
        segments: List[Dict[str, Any]],
        full_text: str,
        title: str = "Transcription",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> bytes:
        """Export transcription as PDF document.

        Args:
            segments: List of segment dicts with start_ms, end_ms, text, confidence_score
            full_text: Full transcription text
            title: Document title
            metadata: Optional metadata dict

        Returns:
            PDF file bytes
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        )

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.75 * inch, bottomMargin=0.75 * inch)

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Title"],
            fontSize=20,
            spaceAfter=12,
        )
        heading_style = ParagraphStyle(
            "CustomHeading",
            parent=styles["Heading2"],
            fontSize=14,
            spaceAfter=8,
            textColor=HexColor("#00FF74"),
        )
        body_style = ParagraphStyle(
            "CustomBody",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
        )
        meta_style = ParagraphStyle(
            "MetaStyle",
            parent=styles["Normal"],
            fontSize=9,
            textColor=HexColor("#666666"),
        )

        story = []
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))

        if metadata:
            story.append(Paragraph("Metadata", heading_style))
            for key, value in metadata.items():
                story.append(Paragraph(f"<b>{key}:</b> {value}", meta_style))
            story.append(Spacer(1, 12))

        story.append(Paragraph("Full Transcription", heading_style))
        story.append(Paragraph(full_text, body_style))
        story.append(Spacer(1, 16))

        if segments:
            story.append(Paragraph("Timestamped Segments", heading_style))

            table_data = [["Time", "Text", "Confidence"]]
            for seg in segments:
                start = _format_timestamp_ms(int(seg.get("start_ms", 0)))
                end = _format_timestamp_ms(int(seg.get("end_ms", 0)))
                confidence = seg.get("confidence_score")
                conf_str = f"{confidence:.1%}" if confidence is not None else "N/A"
                table_data.append([f"{start} → {end}", seg.get("text", ""), conf_str])

            table = Table(table_data, colWidths=[1.5 * inch, 4 * inch, 1 * inch])
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), HexColor("#00FF74")),
                ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#000000")),
                ("FONTSIZE", (0, 0), (-1, 0), 9),
                ("FONTSIZE", (0, 1), (-1, -1), 8),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ALIGN", (0, 0), (-1, 0), "LEFT"),
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#FFFFFF"), HexColor("#F8F8F8")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(table)

        doc.build(story)
        buffer.seek(0)
        return buffer.read()


document_exporter = DocumentExporter()
